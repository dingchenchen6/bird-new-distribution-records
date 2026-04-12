#!/usr/bin/env python3

from pathlib import Path
from shutil import copy2
import re
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


TASK_ROOT = Path("/Users/dingchenchen/Documents/New records/bird-new-distribution-records/tasks/bird_scientific_data_descriptor")
RESULTS_DIR = TASK_ROOT / "results"
FIG_DIR = TASK_ROOT / "figures"
DATA_DIR = TASK_ROOT / "data"
RESULTS_DIR.mkdir(parents=True, exist_ok=True)
FIG_DIR.mkdir(parents=True, exist_ok=True)
DATA_DIR.mkdir(parents=True, exist_ok=True)

SOURCE_MD = RESULTS_DIR / "bird_scientific_data_descriptor_source.md"
OUT_DOCX = RESULTS_DIR / "bird_scientific_data_descriptor_scientific_data.docx"

ORDER_SUMMARY_CSV = Path("/Users/dingchenchen/Documents/New records/bird-new-distribution-records/tasks/bird_identity_synonym_dedup_reanalysis/order_summary_corrected/data/table_order_summary_corrected_numeric.csv")

FIGURE_SOURCES = {
    "FIGURE1": {
        "src": Path("/Users/dingchenchen/Documents/New records/bird-new-distribution-records/tasks/bird_identity_synonym_dedup_reanalysis/sankey_corrected/figures/fig_sankey_order_province_year_corrected.png"),
        "caption": "Figure 1 | Sankey diagram showing the relationships among bird orders, provinces, and publication years represented in the corrected CBNR analytical release.",
        "width": 6.8,
        "filename": "figure1_sankey.png",
    },
    "FIGURE2": {
        "src": Path("/Users/dingchenchen/Documents/New records/bird-new-distribution-records/tasks/bird_identity_synonym_dedup_reanalysis/spatiotemporal_corrected/figures/fig_sp01_province_new_record_count_map.png"),
        "caption": "Figure 2 | Spatial distribution of provincial-level bird new-record counts in the corrected CBNR analytical release.",
        "width": 6.6,
        "filename": "figure2_province_map.png",
    },
    "FIGURE3": {
        "src": Path("/Users/dingchenchen/Documents/New records/bird-new-distribution-records/tasks/bird_identity_synonym_dedup_reanalysis/figures/fig_qa_identity_synonym_duplicate_reanalysis.png"),
        "caption": "Figure 3 | Quality-control overview summarizing the effects of species-identity correction and same-species same-province duplicate removal on the corrected CBNR release.",
        "width": 6.6,
        "filename": "figure3_qa.png",
    },
}

TABLE1_ROWS = [
    ("Species Chinese name", "100"),
    ("Species English name", "100"),
    ("Observation date (year, month, day)", "100"),
    ("Previously reported provinces", "100"),
    ("Province of new record", "100"),
    ("Body length", "100"),
    ("Geographic coordinates (latitude, longitude)", "100"),
    ("Altitude", "100"),
    ("Habitat", "99.3"),
    ("Migratory status", "99.3"),
    ("Identification method", "99.3"),
    ("Reason for new record and significance", "99.3"),
    ("Publication year", "100"),
    ("Authors", "100"),
    ("Journal", "100"),
    ("Article title", "100"),
    ("DOI", "90.7"),
]

TABLE2_ROWS = [
    ("Taxonomic information", "record_id", "Stable row identifier in the corrected event table", "1"),
    ("Taxonomic information", "species_cn", "Chinese common name", "黑翅鸢"),
    ("Taxonomic information", "english_name", "English common name", "Black-winged Kite"),
    ("Taxonomic information", "species", "Accepted scientific name after taxonomic harmonization", "Elanus caeruleus"),
    ("Taxonomic information", "order", "Accepted order name", "Accipitriformes"),
    ("Geographic information", "province", "Province of the validated new record", "Hunan"),
    ("Geographic information", "longitude", "Decimal longitude (WGS84)", "112.9876"),
    ("Geographic information", "latitude", "Decimal latitude (WGS84)", "28.2345"),
    ("Geographic information", "year", "Publication year of the provincial first record", "2023"),
    ("Conservation information", "iucn", "IUCN category standardized from the source and checklist context", "LC"),
    ("Discovery information", "discover_reason", "Primary inferred reason for the new record", "Survey gap or under-sampling"),
    ("Discovery information", "discovery_method", "Observation or detection method reported by the source", "Observe and photograph"),
    ("Audit information", "identity_source", "Primary source used for species-identity correction", "true_mismatch"),
    ("Audit information", "identity_change_flag", "Whether the accepted binomial changed during correction", "TRUE"),
    ("Audit information", "duplicate_group_size", "Number of rows in the same species-province duplicate group", "2"),
]


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        elem = tcBorders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            tcBorders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "6")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "B0B0B0")


def style_table(table, header_fill="D9E2F3"):
    for r, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)
            set_cell_border(cell)
        if r == 0:
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), header_fill)
                tcPr.append(shd)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.size = Pt(9)


def add_table_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.name = "Arial"
    run.font.size = Pt(9)


def add_figure(doc, key):
    spec = FIGURE_SOURCES[key]
    target = FIG_DIR / spec["filename"]
    if not target.exists():
        copy2(spec["src"], target)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(target), width=Inches(spec["width"]))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(spec["caption"])
    run.italic = True
    run.font.name = "Arial"
    run.font.size = Pt(9)


def add_table1(doc):
    add_table_caption(doc, "Table 1 | Field-level accuracy in the AI-assisted extraction calibration set (100 articles).")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Information category"
    hdr[1].text = "Accuracy (%)"
    for row in TABLE1_ROWS:
        cells = table.add_row().cells
        cells[0].text = row[0]
        cells[1].text = row[1]
    style_table(table)
    doc.add_paragraph()


def add_table2(doc):
    add_table_caption(doc, "Table 2 | Representative metadata fields in the corrected CBNR analytical release.")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Unit", "Variable name", "Description", "Example"]):
        hdr[i].text = h
    for row in TABLE2_ROWS:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    style_table(table)
    doc.add_paragraph()


def add_table3(doc):
    df = pd.read_csv(ORDER_SUMMARY_CSV)
    df = df[df["order"] != "Total"].copy()
    df["Proportion of newly recorded bird species"] = (df["prop_new_species_all"] * 100).map(lambda x: f"{x:.1f}%")
    df["Proportion to total species in the order"] = (df["prop_to_total_species_order"] * 100).map(lambda x: f"{x:.1f}%")
    show = df[[
        "order",
        "n_new_species",
        "n_papers",
        "Proportion of newly recorded bird species",
        "Proportion to total species in the order",
    ]].rename(columns={
        "order": "Order",
        "n_new_species": "Number of newly recorded bird species",
        "n_papers": "Number of papers",
    })
    add_table_caption(doc, "Table 3 | Summary statistics of newly recorded bird species by order in the corrected CBNR analytical release.")
    table = doc.add_table(rows=1, cols=len(show.columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(show.columns):
        hdr[i].text = str(h)
    for _, row in show.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row.tolist()):
            cells[i].text = str(value)
    style_table(table)
    doc.add_paragraph()


def add_heading(doc, text, level):
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    run.font.name = "Arial"
    if level == 1:
        run.font.size = Pt(18)
        run.bold = True
    elif level == 2:
        run.font.size = Pt(14)
        run.bold = True
    else:
        run.font.size = Pt(12)
        run.bold = True


def add_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    for part in re.split(r"(\*[^*]+\*)", text):
        if not part:
            continue
        if part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            run = p.add_run(part)
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
    return p


def build_doc():
    text = SOURCE_MD.read_text(encoding="utf-8")
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)

    lines = text.splitlines()
    paragraph_buffer = []

    def flush_buffer():
        nonlocal paragraph_buffer
        if paragraph_buffer:
            para = " ".join(x.strip() for x in paragraph_buffer).strip()
            if para:
                add_paragraph(doc, para)
            paragraph_buffer = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            flush_buffer()
            continue
        if stripped == "[[TABLE1]]":
            flush_buffer()
            add_table1(doc)
            continue
        if stripped == "[[TABLE2]]":
            flush_buffer()
            add_table2(doc)
            continue
        if stripped == "[[TABLE3]]":
            flush_buffer()
            add_table3(doc)
            continue
        if stripped in ("[[FIGURE1]]", "[[FIGURE2]]", "[[FIGURE3]]"):
            flush_buffer()
            add_figure(doc, stripped.strip("[]"))
            continue
        if stripped.startswith("# "):
            flush_buffer()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped[2:])
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(18)
            doc.add_paragraph()
            continue
        if stripped.startswith("## "):
            flush_buffer()
            add_heading(doc, stripped[3:], 2)
            continue
        if stripped.startswith("### "):
            flush_buffer()
            add_heading(doc, stripped[4:], 3)
            continue
        paragraph_buffer.append(stripped)

    flush_buffer()
    doc.save(str(OUT_DOCX))


if __name__ == "__main__":
    build_doc()
    print(f"Wrote {OUT_DOCX}")
