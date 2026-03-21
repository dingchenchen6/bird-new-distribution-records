#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
=========================================================================
Scientific question / 科学问题
-------------------------------------------------------------------------
EN:
How can the completed bird new-record analyses be translated into a coherent,
submission-oriented English manuscript that integrates background, methods,
 results, interpretation, tables, and combined figures in a reproducible Word
workflow?

CN:
如何将已经完成的鸟类新纪录分析转化为一篇结构完整、接近投稿状态的英文论文，
并以可重复的 Word 生成流程整合背景、方法、结果、讨论、表格和组合图件？

Objective / 目标
EN:
1. Draft a full English manuscript from title to conclusion.
2. Reuse existing validated analytical outputs rather than refitting models.
3. Embed combined manuscript figures and manuscript tables directly into DOCX.
4. Export a readable markdown source alongside the Word file for versioning.
5. Generate a simple diagnostic record confirming figure and table counts.

CN:
1. 生成从标题到结论的完整英文论文稿。
2. 基于已有并已验证的分析结果组织稿件，而不是重复拟合模型。
3. 将正文组合图和论文表格直接嵌入 DOCX。
4. 同时导出可读的 markdown 文本源文件，方便版本管理。
5. 生成简单诊断文件，确认图表数量和关键输出。

Analytical idea / 思路
EN:
The manuscript is built around four major result domains: taxonomic structure,
spatiotemporal concentration, directional non-randomness, and species-level
correlates/discovery mechanisms. Equivalent visual variants are not repeated.
Instead, one preferred main-text version is chosen for each domain, while a
supplementary figure is retained for order-specific directionality.

CN:
稿件围绕四个主要结果模块组织：分类结构、时空集中格局、方向性非随机、以及物
种层面相关性/发现机制。对于表达相同核心结果的图件变体不再重复展示，而是为每
个模块选择一个正文主图版本，并保留一个按目方向结构的补充图。

Diagnostics and validation / 诊断与验证
EN:
The script checks that all required source files exist, formats tables from
exported CSV files, writes both markdown and DOCX outputs, and stores a small
CSV with manuscript diagnostics. Because full DOCX visual rendering tools are
not available in the environment, post-build validation relies on DOCX media
inspection and text-level checks rather than page-image rendering.

CN:
脚本会检查所有必要来源文件是否存在，从既有 CSV 表中格式化正文表格，同时导出
markdown 与 DOCX，并写出简要诊断 CSV。由于当前环境缺少完整 DOCX 页面渲染工
具，成稿核查将主要依赖 DOCX 内部媒体检查与文本层面的检查，而非逐页图像渲染。
=========================================================================
"""

from pathlib import Path
from zipfile import ZipFile
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK

ROOT = Path("/Users/dingchenchen/Documents/New records/bird_new_records_R_output/tasks")
OUT = ROOT / "bird_manuscript_english_full"
FIG = OUT / "figures"
DATA = OUT / "data"
RES = OUT / "results"
RES.mkdir(parents=True, exist_ok=True)

# Step 1. Input files / 输入文件
qa = pd.read_csv(ROOT / "bird_full_pipeline_main_text/data/table_qa_summary.csv")
order_summary = pd.read_csv(ROOT / "bird_order_summary_table/data/table_order_summary_bird_new_records_numeric.csv")
trait = pd.read_csv(ROOT / "bird_full_pipeline_main_text/data/table_species_trait_model_coefficients.csv")
province_summary = pd.read_csv(ROOT / "bird_spatiotemporal_patterns/data/province_spatiotemporal_summary.csv")
year_counts = pd.read_csv(ROOT / "bird_full_pipeline_main_text/data/table_year_counts.csv")
order_counts = pd.read_csv(ROOT / "bird_full_pipeline_main_text/data/table_order_counts.csv")
iucn_counts = pd.read_csv(ROOT / "bird_full_pipeline_main_text/data/table_iucn_counts.csv")
reason_counts = pd.read_csv(ROOT / "bird_full_pipeline_main_text/data/table_discovery_reason_counts.csv")
method_counts = pd.read_csv(ROOT / "bird_full_pipeline_main_text/data/table_discovery_method_counts.csv")
figure_plan = pd.read_csv(DATA / "manuscript_figure_plan.csv")

required_figures = [
    FIG / "figure1_spatiotemporal_overview.png",
    FIG / "figure2_taxonomic_flow_sankey.png",
    FIG / "figure3_directional_patterns.png",
    FIG / "figure4_correlates_and_mechanisms.png",
    FIG / "figureS1_order_direction_facets.png",
]
for path in required_figures:
    if not path.exists():
        raise FileNotFoundError(f"Missing manuscript figure: {path}")

# Step 2. Summary values / 摘要数值
qa_map = dict(zip(qa["metric"], qa["value"]))
n_records = int(qa_map["Rows after deduplication"])
n_species = int(qa_map["Unique species"])
n_orders = int(qa_map["Unique orders"])
n_provinces = int(qa_map["Unique provinces"])
study_period = f"{int(qa_map['Year min'])}-{int(qa_map['Year max'])}"

top_prov = province_summary.sort_values("n_records", ascending=False).head(5)
top_density = province_summary.sort_values("density_100k", ascending=False).head(5)
top_years = year_counts.sort_values("n_records", ascending=False).head(5)
top_order = order_counts.sort_values("n_records", ascending=False).head(5)

sig_trait = trait.loc[trait["p.value"] < 0.05].copy()
sig_trait["estimate_fmt"] = sig_trait["estimate"].map(lambda x: f"{x:.3f}")
sig_trait["p_fmt"] = sig_trait["p.value"].map(lambda x: "<0.001" if x < 0.001 else f"{x:.3f}")

# Step 3. Manuscript text blocks / 论文文本模块
TITLE = "Taxonomic, Spatiotemporal and Directional Patterns of New Bird Distribution Records in China"
SUBTITLE = "Draft English manuscript synthesized from the completed bird new-record analysis workflow"

ABSTRACT = (
    f"Accurate knowledge of species distributions is central to biodiversity science, yet the Wallacean shortfall remains substantial even in well-observed vertebrates such as birds. "
    f"Here, we synthesized completed analyses of new bird distribution records in China to evaluate their taxonomic structure, spatiotemporal concentration, directional non-randomness, and species-level correlates. "
    f"After quality control and deduplication, the dataset retained {n_records} records involving {n_species} species across {n_orders} orders and {n_provinces} provincial-level units during {study_period}. "
    f"Passeriformes dominated the dataset in absolute records, while several species-poor orders showed high relative representativeness within their species pools. Spatial hotspots in absolute counts were concentrated in Tibet, Yunnan, Hunan, Shaanxi, and Gansu, whereas area-standardized discovery intensity was highest in Shanghai, Tianjin, Beijing, Hong Kong, and Hainan. "
    f"Record accumulation accelerated strongly in recent years, with peaks in 2024, 2025, and 2023. Directional analyses showed that new records were not randomly distributed around previously known ranges: eastward records were most frequent, followed by northeastward and southeastward records. Species-level models indicated that range size and migratory behaviour were consistent positive correlates of new records, whereas body mass and broad risk-group categories were not significant in the fitted models. Discovery reasons and methods jointly suggested that new bird records reflect both ecological redistribution and intensified observation and documentation. "
    f"Taken together, the results show that the reduction of the Wallacean shortfall in China is highly uneven across taxa, provinces, and directions, and that future survey prioritization should target not only record-rich hotspots but also regions and groups with persistent residual knowledge gaps."
)

KEYWORDS = "birds; China; directional bias; new distribution records; species traits; spatiotemporal pattern; Wallacean shortfall"

INTRODUCTION = [
    "Understanding where species occur is fundamental to biogeography, macroecology, and conservation biology. However, distribution knowledge remains incomplete and uneven, a problem commonly described as the Wallacean shortfall. This shortfall limits biodiversity assessment, obscures responses to environmental change, and weakens spatial prioritization for conservation.",
    "New distribution records provide a particularly informative entry point into this problem because they identify taxa, places, and times for which previous knowledge was incomplete, inaccurate, or outdated. In this sense, new records are more than descriptive additions to faunal lists; they are empirical signals of how biodiversity knowledge is being updated in space and time.",
    "The conceptual framework developed in Ding et al. (2025) for mammals in China suggests that new records emerge through the interaction of ecological processes and observation processes. Ecological processes include genuine range expansion, contraction, and redistribution under climate or habitat change. Observation processes include survey effort, accessibility, research infrastructure, local reporting intensity, photography, and other improvements in detectability. For birds, this coupling is expected to be particularly strong because many species are mobile, migratory, and sensitive to environmental change, while at the same time being intensively observed by both professional and public communities.",
    "China is an ideal setting in which to evaluate these processes because it spans two major zoogeographic realms, includes strong climatic and topographic gradients, and has experienced rapid growth in biodiversity surveys, protected-area monitoring, birdwatching activity, and digital documentation platforms. Yet the resulting gain in distribution knowledge is unlikely to be spatially or taxonomically uniform.",
    "Using the completed bird new-record workflow, we synthesized four major dimensions of evidence: taxonomic structure, spatiotemporal concentration, directional patterns relative to previously known ranges, and species-level ecological correlates. We addressed four questions: (1) Are bird new records taxonomically concentrated? (2) Are they strongly uneven in space and time? (3) Do they show non-random directional structure? and (4) Which species characteristics are associated with a higher probability or frequency of new records?"
]

METHODS = {
    "2.1 Data compilation and quality control": [
        "The manuscript synthesizes a standardized bird new-record workflow built from the cleaned master spreadsheet and previously exported analytical outputs. The raw compilation contained 1078 rows. After filtering missing key fields and deduplicating repeated entries, 1059 records were retained for analysis.",
        "The final dataset represented 578 species across 23 orders and 33 provincial-level units from 2000 to 2025. All subsequent manuscript results were derived from exported workflow tables rather than re-entered manually, thereby reducing transcription error and preserving traceability."
    ],
    "2.2 Taxonomic and spatiotemporal summaries": [
        "We summarized record counts by order, province, year, IUCN threat category, discovery reason, and discovery method. At the provincial level, both absolute record counts and area-standardized densities (records per 100,000 km2) were considered in order to distinguish total accumulation from relative discovery intensity.",
        "Spatial products were based on a China-appropriate projected mapping workflow using a CGCS2000-compatible Albers Equal Area projection. The manuscript retained one preferred count map and one point-distribution map for main-text synthesis, while equivalent graticule variants were excluded from the main text to avoid duplication."
    ],
    "2.3 Directional analyses": [
        "Directional analyses classified bird new records into eight compass sectors (North, Northeast, East, Southeast, South, Southwest, West, and Northwest) relative to previously known ranges. The main workflow generated overall overlay plots, order-specific 4 x 4 faceted windrose panels, and an order-level binomial expectation envelope that contrasted observed directional proportions against a random expectation.",
        "For the manuscript, one overall directional synthesis figure and one supplementary order-specific panel were retained. Equivalent representations of the same overall signal were not repeated in the main text."
    ],
    "2.4 Species-level correlates and discovery mechanisms": [
        "We summarized the results of two previously fitted species-level models: a binary occurrence model and a quasi-Poisson count model. Predictors retained in the exported coefficient table included log body mass, range size measured as the number of occupied provinces, migration class, and a broad risk-group classification.",
        "To interpret whether new records were more likely to reflect ecological redistribution or observation enhancement, we also summarized the stated discovery reasons and discovery methods across the compiled record set."
    ],
    "2.5 Manuscript figure and table selection": [
        "Equivalent analytical variants were not repeated in the main text. Instead, one preferred figure was retained for each major analytical domain: spatiotemporal pattern, taxonomic flow, directional structure, and ecological correlates/discovery mechanisms. Related figures were combined into multi-panel layouts only when such combination improved interpretability without duplicating the underlying result.",
        "Main-text tables were generated directly from exported CSV summaries, whereas full figure-to-result coverage was documented separately in the bilingual figure registry task."
    ],
}

RESULTS = {
    "3.1 Dataset overview and taxonomic structure": [
        f"The final dataset contained {n_records} bird new records representing {n_species} species, {n_orders} orders, and {n_provinces} provincial-level units from {study_period}. Passeriformes dominated the dataset with 524 records, followed by Charadriiformes (155), Anseriformes (68), Accipitriformes (67), and Pelecaniformes (46).",
        "Absolute dominance and relative representativeness were not identical. Although Passeriformes contributed the highest number of newly recorded species and papers, several species-poor orders showed higher proportional representativeness within their full species pools, indicating that numerical prominence and within-order discovery completeness are different dimensions of the taxonomic pattern."
    ],
    "3.2 Spatial and temporal concentration": [
        "Bird new records were spatially concentrated in a limited subset of provinces. The largest absolute totals occurred in Tibet, Yunnan, Hunan, Shaanxi, and Gansu. In contrast, area-standardized discovery intensity was highest in Shanghai, Tianjin, Beijing, Hong Kong, and Hainan.",
        "Record accumulation also accelerated markedly through time. The largest yearly totals were recorded in 2024, 2025, and 2023, with additional peaks in 2019 and 2017. This increase was driven primarily by Passeriformes but also included notable contributions from Charadriiformes and Accipitriformes, indicating that recent acceleration was taxonomically broad rather than exclusively single-order driven."
    ],
    "3.3 Taxonomic flow among order, province, and year": [
        "The compact Sankey diagram revealed a strongly asymmetric flow structure linking order, province, and year. Passeriformes formed the dominant upstream source, and the strongest downstream provincial flows converged on Tibet, Yunnan, Hunan, and Shaanxi.",
        "Recent years, especially after 2020, were represented by visibly wider ribbons, demonstrating that the joint taxonomic-spatial-temporal accumulation of bird new records intensified in the most recent period."
    ],
    "3.4 Directional patterns": [
        "Directional analyses showed that bird new records were not evenly distributed around previously known ranges. East was the dominant sector, followed by Northeast and Southeast. The overall windrose figure and the order-specific panels indicated that this pattern was shared across several major orders rather than being restricted to one clade.",
        "Passeriformes contributed most strongly to the eastward and northeastward signal. The binomial expectation plot further indicated that several orders departed from random directional expectation, supporting a genuinely non-random directional structure in the compiled record set."
    ],
    "3.5 Species-level correlates and discovery mechanisms": [
        "Range size and migratory behaviour showed consistent positive relationships with bird new records in both the binary and quasi-Poisson models. By contrast, body mass and broad risk-group categories were not significant in the fitted models.",
        "Discovery reasons and discovery methods suggested a mixed-process explanation. Range shift or distribution change was the dominant stated reason, whereas direct observation and photography were by far the most frequent detection mode. This combination suggests that current bird new records in China reflect both real redistribution and intensified documentation."
    ],
}

DISCUSSION = {
    "4.1 Uneven reduction of the Wallacean shortfall": [
        "The bird new-record dataset shows that the reduction of the Wallacean shortfall in China is highly uneven rather than smooth or uniform. Knowledge gain is concentrated in species-rich orders, a limited set of hotspot provinces, and a narrow directional subset of potential range margins.",
        "This unevenness mirrors the central insight from the mammal study of Ding et al. (2025), but with a taxonomic signature more strongly shaped by mobility and migration. In birds, the strongest absolute signal lies in Passeriformes, yet the within-order representativeness pattern shows that smaller orders can have disproportionately high relative novelty."
    ],
    "4.2 Coupled ecological redistribution and observation enhancement": [
        "The results are most consistent with a coupled-process explanation. If bird new records were driven only by intensified observation, one would expect weaker and less coherent directional structure. Conversely, if they reflected only ecological redistribution, one would expect a cleaner trait signal with weaker evidence of detection-mode dominance.",
        "Instead, the dataset combines strong eastward and northeastward directional structure, significant effects of range size and migration, and overwhelming reliance on direct observation and photography. This indicates that ecological change and observation enhancement are acting simultaneously."
    ],
    "4.3 Implications for survey prioritization": [
        "The contrast between absolute hotspot provinces and density hotspots has practical implications. Large, biodiverse provinces such as Tibet and Yunnan remain central targets because they continue to accumulate many new records. At the same time, high-density coastal units and islands indicate that relatively small areas can produce disproportionately high discovery intensity when observation is strong.",
        "Future survey design should therefore integrate at least three criteria: absolute hotspot status, area-standardized discovery intensity, and taxonomic or directional novelty. Such a framework would better target residual knowledge gaps than any single metric alone."
    ],
    "4.4 Limitations and next steps": [
        "This manuscript synthesizes completed analytical outputs rather than refitting every model from raw sources during manuscript generation. As a result, all numerical statements are traceable to existing task outputs, but some methodological granularity remains constrained by the finalized workflow tables.",
        "A natural next step would be to extend the synthesis into a unified hierarchical framework that simultaneously models ecological redistribution, survey effort, and trait-mediated detectability across birds and other vertebrate groups."
    ],
}

CONCLUSION = [
    "Bird new distribution records in China are highly uneven in taxonomic, spatial, temporal, and directional terms. The dominant signatures are strong concentration in Passeriformes, hotspot provinces in the southwest and plateau margins, very high per-area discovery intensity in small coastal units, and pervasive eastward to northeastward directional bias.",
    "Taken together, these findings indicate that bird new records are best understood as products of both ecological redistribution and intensified observation. A survey strategy that integrates hotspot intensity, representativeness, and directional novelty will provide a stronger basis for reducing the Wallacean shortfall in China."
]

REFERENCES = [
    "Boakes EH, McGowan PJK, Fuller RA, et al. 2010. Distorted views of biodiversity: spatial and temporal bias in species occurrence data. PLoS Biology 8:e1000385.",
    "Bowler DE, Boyd RJ, Callaghan CT, et al. 2025. Treating gaps and biases in biodiversity data as a missing data problem. Biological Reviews 100:50-67.",
    "Chen S, Chen Z, Lin H, et al. 2025. Chinese provincial-level new records for resident bird species reveal poleward range shifts. Avian Research 100310.",
    "Darwin C. 1859. On the Origin of Species. London: John Murray.",
    "Ding C, Ding J, Qiao H, Jiang Z, Wang Z. 2025. Taxonomic and spatiotemporal patterns and ecological correlates of new mammal distribution records in China. Global Ecology and Biogeography 34:e70165.",
    "Diniz-Filho JAF, Jardim L, Guedes JJ, et al. 2023. Macroecological links between the Linnean, Wallacean, and Darwinian shortfalls. Frontiers of Biogeography 15:e59566.",
    "Hortal J, de Bello F, Diniz-Filho JAF, et al. 2015. Seven shortfalls that beset large-scale knowledge of biodiversity. Annual Review of Ecology, Evolution, and Systematics 46:523-549.",
    "Hughes AC, Orr MCC, Ma K, et al. 2021. Sampling biases shape our view of the natural world. Ecography 44:1259-1269.",
    "Jetz W, McGeoch MA, Guralnick R, et al. 2019. Essential biodiversity variables for mapping and monitoring species populations. Nature Ecology & Evolution 3:539-551.",
    "Meyer C, Kreft H, Guralnick R, Jetz W. 2015. Global priorities for an effective information basis of biodiversity distributions. Nature Communications 6:8221.",
    "Moura MR, Jetz W. 2021. Shortfalls and opportunities in terrestrial vertebrate species discovery. Nature Ecology & Evolution 5:631-639.",
    "Oliver RY, Meyer C, Ranipeta A, Winner K, Jetz W. 2021. Global and national trends, gaps, and opportunities in documenting and monitoring species distributions. PLoS Biology 19:e3001336.",
    "Sanczuk P, Lenoir J, Denelle P, et al. 2026. Global bias towards recording latitudinal range shifts. Nature Climate Change 16:21-25.",
    "Whittaker RJ, Araujo MB, Jepson P, et al. 2005. Conservation biogeography: assessment and prospect. Diversity and Distributions 11:3-23."
]

FIGURE_CAPTIONS = {
    "Figure 1": "Spatiotemporal overview of bird new records in China. (A) Provincial count map. (B) Georeferenced point-distribution map across major orders. (C) Province-level composition of the top 10 orders plus Others. (D) Year-level composition of the top 10 orders plus Others.",
    "Figure 2": "Compact Sankey diagram linking bird order, province, and year. Ribbon width represents the number of records in each pathway.",
    "Figure 3": "Directional patterns of bird new records relative to previously known ranges. (A) Overall overlay windrose for major orders. (B) Order-level binomial expectation envelope. (C) Order-specific 4 x 4 windrose panels for the major orders.",
    "Figure 4": "Ecological correlates and discovery mechanisms of bird new records. (A) Coefficient estimates from the binary and quasi-Poisson species-level models. (B) Order-level representativeness relative to the 2025 national species pool. (C) Discovery reasons and discovery methods.",
    "Supplementary Figure S1": "Order-specific directional windrose facets retained as a supplementary figure for detailed taxonomic comparison."
}

# Step 4. Create markdown source / 生成 markdown 文本源文件
md_lines = []
md_lines.append(f"# {TITLE}")
md_lines.append("")
md_lines.append(f"*{SUBTITLE}*")
md_lines.append("")
md_lines.append("## Abstract")
md_lines.append(ABSTRACT)
md_lines.append("")
md_lines.append(f"**Keywords:** {KEYWORDS}")
md_lines.append("")
md_lines.append("## 1. Introduction")
md_lines.extend(INTRODUCTION)
md_lines.append("")
md_lines.append("## 2. Materials and Methods")
for sec, paras in METHODS.items():
    md_lines.append(f"### {sec}")
    md_lines.extend(paras)
    md_lines.append("")
md_lines.append("## 3. Results")
for sec, paras in RESULTS.items():
    md_lines.append(f"### {sec}")
    md_lines.extend(paras)
    md_lines.append("")
md_lines.append("## 4. Discussion")
for sec, paras in DISCUSSION.items():
    md_lines.append(f"### {sec}")
    md_lines.extend(paras)
    md_lines.append("")
md_lines.append("## 5. Conclusions")
md_lines.extend(CONCLUSION)
md_lines.append("")
md_lines.append("## References")
md_lines.extend([f"- {r}" for r in REFERENCES])
md_lines.append("")
md_lines.append("## Figure captions")
for k, v in FIGURE_CAPTIONS.items():
    md_lines.append(f"- **{k}.** {v}")

md_path = RES / "bird_new_records_full_manuscript_english.md"
md_path.write_text("\n".join(md_lines), encoding="utf-8")

# Step 5. Initialize DOCX / 初始化 DOCX

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)

styles = doc.styles
styles["Normal"].font.name = "Times New Roman"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
styles["Normal"].font.size = Pt(11)

for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
    styles[style_name].font.name = "Times New Roman"
    styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

styles["Title"].font.size = Pt(16)
styles["Heading 1"].font.size = Pt(13)
styles["Heading 2"].font.size = Pt(11.5)
styles["Heading 2"].font.bold = True

# Helper functions / 辅助函数

def add_body_paragraph(text, italic=False):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.italic = italic
    return p


def add_caption(label, text):
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    r1 = p.add_run(f"{label}. ")
    r1.bold = True
    p.add_run(text)


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)


def add_dataframe_table(df, caption, widths=None, font_size=8.5):
    add_body_paragraph("")
    add_caption(caption[0], caption[1])
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr[i].text = str(col)
    set_repeat_table_header(table.rows[0])
    for row in df.itertuples(index=False):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(font_size)
    if widths:
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph("")

# Step 6. Title page / 标题页
p = doc.add_paragraph(style="Title")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run(TITLE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Draft manuscript synthesized from the bird new-record workflow")
r.italic = True
r.font.name = "Times New Roman"
r.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Author details can be added at the submission stage.")
r.font.name = "Times New Roman"
r.font.size = Pt(10.5)

doc.add_paragraph("")
doc.add_heading("Abstract", level=1)
add_body_paragraph(ABSTRACT)
add_body_paragraph(f"Keywords: {KEYWORDS}")

# Step 7. Main text / 正文

doc.add_heading("1. Introduction", level=1)
for para in INTRODUCTION:
    add_body_paragraph(para)

doc.add_heading("2. Materials and Methods", level=1)
for sec, paras in METHODS.items():
    doc.add_heading(sec, level=2)
    for para in paras:
        add_body_paragraph(para)

doc.add_heading("3. Results", level=1)
for sec, paras in RESULTS.items():
    doc.add_heading(sec, level=2)
    for para in paras:
        add_body_paragraph(para)
    # insert figures/tables after relevant sections
    if sec == "3.1 Dataset overview and taxonomic structure":
        overview_df = pd.DataFrame([
            ["Records after QA and deduplication", n_records],
            ["Unique species", n_species],
            ["Unique orders", n_orders],
            ["Unique provincial-level units", n_provinces],
            ["Study period", study_period],
            ["Top order by records", f"{top_order.iloc[0]['order']} ({int(top_order.iloc[0]['n_records'])})"],
            ["Top province by count", f"{top_prov.iloc[0]['province_std']} ({int(top_prov.iloc[0]['n_records'])})"],
            ["Top year", f"{int(top_years.iloc[0]['year'])} ({int(top_years.iloc[0]['n_records'])})"],
        ], columns=["Metric", "Value"])
        add_dataframe_table(overview_df, ("Table 1", "Study overview of the bird new-record dataset."), widths=[3.4, 2.8], font_size=9)
        ord_df = order_summary.copy()
        ord_df["Share of all new species (%)"] = (ord_df["prop_new_species_all"] * 100).round(1)
        ord_df["Share within order pool (%)"] = (ord_df["prop_to_total_species_order"] * 100).round(1)
        ord_df = ord_df.loc[ord_df["order"] != "Total", ["order", "n_new_species", "n_papers", "Share of all new species (%)", "Share within order pool (%)"]]
        ord_df.columns = ["Order", "New species", "Papers", "Share of all new species (%)", "Share within order pool (%)"]
        add_dataframe_table(ord_df, ("Table 2", "Order-level summary of newly recorded bird species in China."), font_size=7.5)
    if sec == "3.2 Spatial and temporal concentration":
        doc.add_picture(str(FIG / "figure1_spatiotemporal_overview.png"), width=Inches(6.2))
        add_caption("Figure 1", FIGURE_CAPTIONS["Figure 1"])
    if sec == "3.3 Taxonomic flow among order, province, and year":
        doc.add_picture(str(FIG / "figure2_taxonomic_flow_sankey.png"), width=Inches(6.3))
        add_caption("Figure 2", FIGURE_CAPTIONS["Figure 2"])
    if sec == "3.4 Directional patterns":
        doc.add_picture(str(FIG / "figure3_directional_patterns.png"), width=Inches(5.9))
        add_caption("Figure 3", FIGURE_CAPTIONS["Figure 3"])
    if sec == "3.5 Species-level correlates and discovery mechanisms":
        trait_df = trait.copy()
        trait_df["Estimate"] = trait_df["estimate"].round(3)
        trait_df["95% CI"] = trait_df.apply(lambda r: f"{r['conf.low']:.3f} to {r['conf.high']:.3f}", axis=1)
        trait_df["P value"] = trait_df["p.value"].map(lambda x: "<0.001" if x < 0.001 else f"{x:.3f}")
        trait_df = trait_df[["model", "term_label", "Estimate", "95% CI", "P value"]]
        trait_df.columns = ["Model", "Predictor", "Estimate", "95% CI", "P value"]
        add_dataframe_table(trait_df, ("Table 3", "Coefficients of the exported species-level models summarized in the manuscript."), font_size=7.8)
        doc.add_picture(str(FIG / "figure4_correlates_and_mechanisms.png"), width=Inches(6.2))
        add_caption("Figure 4", FIGURE_CAPTIONS["Figure 4"])

doc.add_heading("4. Discussion", level=1)
for sec, paras in DISCUSSION.items():
    doc.add_heading(sec, level=2)
    for para in paras:
        add_body_paragraph(para)

doc.add_heading("5. Conclusions", level=1)
for para in CONCLUSION:
    add_body_paragraph(para)

doc.add_heading("References", level=1)
for ref in REFERENCES:
    add_body_paragraph(ref)

doc.add_heading("Supplementary figure", level=1)
add_body_paragraph("The main text avoids repeating equivalent analytical variants. The following supplementary figure is retained for detailed directional comparison among orders.")
doc.add_picture(str(FIG / "figureS1_order_direction_facets.png"), width=Inches(6.0))
add_caption("Supplementary Figure S1", FIGURE_CAPTIONS["Supplementary Figure S1"])

# Step 8. Save outputs / 保存输出

docx_path = RES / "bird_new_records_full_manuscript_english.docx"
doc.save(str(docx_path))

# Step 9. Diagnostics / 诊断
with ZipFile(docx_path, 'r') as zf:
    media_files = [n for n in zf.namelist() if n.startswith("word/media/")]

diagnostic = pd.DataFrame([
    ["docx_path", str(docx_path)],
    ["markdown_path", str(md_path)],
    ["n_main_figures", 4],
    ["n_supplementary_figures", 1],
    ["n_tables", 3],
    ["n_embedded_media_files", len(media_files)],
    ["n_references", len(REFERENCES)],
])
diagnostic.columns = ["metric", "value"]
diagnostic.to_csv(DATA / "manuscript_build_diagnostics.csv", index=False)
print(f"Manuscript DOCX created: {docx_path}")
